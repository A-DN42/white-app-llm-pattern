"""
File processing service for different document formats.
Supports CSV, XLSX, PDF, DOCX, and TXT files.
"""

import io
import logging
from typing import List, Dict, Any, Union
from pathlib import Path

import pandas as pd
import PyPDF2
from docx import Document as DocxDocument

# Configure logging
logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Service class for processing different file formats.
    Extracts text content from various document types.
    """
    
    SUPPORTED_EXTENSIONS = {'.csv', '.xlsx', '.xls', '.pdf', '.docx', '.txt'}
    
    def __init__(self):
        logger.info("File processor initialized")
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if the file extension is supported."""
        return Path(filename).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def process_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process a file and extract text content.
        
        Args:
            file_content: The file content as bytes
            filename: The original filename
            
        Returns:
            Dict with success status, extracted text, and metadata
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if not self.is_supported_file(filename):
                return {
                    "success": False,
                    "message": f"Unsupported file type: {file_extension}. Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
                }
            
            # Route to appropriate processor
            if file_extension == '.csv':
                return self._process_csv(file_content, filename)
            elif file_extension in ['.xlsx', '.xls']:
                return self._process_excel(file_content, filename)
            elif file_extension == '.pdf':
                return self._process_pdf(file_content, filename)
            elif file_extension == '.docx':
                return self._process_docx(file_content, filename)
            elif file_extension == '.txt':
                return self._process_txt(file_content, filename)
            else:
                return {
                    "success": False,
                    "message": f"Handler not implemented for {file_extension}"
                }
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing file: {str(e)}"
            }
    
    def _process_csv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV files."""
        try:
            # Read CSV into DataFrame
            df = pd.read_csv(io.BytesIO(file_content))
            
            # Convert to text representation
            text_content = []
            
            # Add column headers
            headers = ", ".join(df.columns.tolist())
            text_content.append(f"CSV File: {filename}")
            text_content.append(f"Columns: {headers}")
            text_content.append("")
            
            # Add each row as text
            for index, row in df.iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                text_content.append(f"Row {index + 1}: " + ", ".join(row_text))
            
            extracted_text = "\n".join(text_content)
            
            return {
                "success": True,
                "text": extracted_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "CSV",
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist()
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing CSV file: {str(e)}"
            }
    
    def _process_excel(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process Excel files (XLSX, XLS)."""
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            text_content = []
            total_rows = 0
            
            text_content.append(f"Excel File: {filename}")
            text_content.append(f"Sheets: {', '.join(excel_file.sheet_names)}")
            text_content.append("")
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                text_content.append(f"Sheet: {sheet_name}")
                headers = ", ".join(df.columns.tolist())
                text_content.append(f"Columns: {headers}")
                text_content.append("")
                
                # Add each row as text
                for index, row in df.iterrows():
                    row_text = []
                    for col, value in row.items():
                        if pd.notna(value):
                            row_text.append(f"{col}: {value}")
                    text_content.append(f"Row {index + 1}: " + ", ".join(row_text))
                
                text_content.append("")
                total_rows += len(df)
            
            extracted_text = "\n".join(text_content)
            
            return {
                "success": True,
                "text": extracted_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "Excel",
                    "sheets": len(excel_file.sheet_names),
                    "sheet_names": excel_file.sheet_names,
                    "total_rows": total_rows
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing Excel file: {str(e)}"
            }
    
    def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF files."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            
            text_content.append(f"PDF File: {filename}")
            text_content.append(f"Pages: {len(pdf_reader.pages)}")
            text_content.append("")
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"Page {page_num}:")
                    text_content.append(page_text.strip())
                    text_content.append("")
            
            extracted_text = "\n".join(text_content)
            
            return {
                "success": True,
                "text": extracted_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "PDF",
                    "pages": len(pdf_reader.pages)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing PDF file: {str(e)}"
            }
    
    def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX files."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_content = []
            
            text_content.append(f"Word Document: {filename}")
            text_content.append("")
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                text_content.append("\nTable:")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                text_content.append("")
            
            extracted_text = "\n".join(text_content)
            
            return {
                "success": True,
                "text": extracted_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "Word Document",
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing DOCX file: {str(e)}"
            }
    
    def _process_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process TXT files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                return {
                    "success": False,
                    "message": "Could not decode text file with supported encodings"
                }
            
            return {
                "success": True,
                "text": f"Text File: {filename}\n\n{text}",
                "metadata": {
                    "filename": filename,
                    "file_type": "Text",
                    "characters": len(text),
                    "lines": len(text.splitlines())
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT file {filename}: {e}")
            return {
                "success": False,
                "message": f"Error processing TXT file: {str(e)}"
            }


# Global file processor instance
_file_processor: Union[FileProcessor, None] = None


def get_file_processor() -> FileProcessor:
    """Get or create the global file processor instance."""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor()
    return _file_processor 